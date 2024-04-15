from django.shortcuts import render
import re
from django.http import HttpResponse
from .models import CV
from .forms import CVUploadForm
import openpyxl
import fitz  # PyMuPDF
from docx import Document
from django.conf import settings
from django.core.files.storage import FileSystemStorage
import os


def extract_info(text):
    email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    phone_pattern = r'(\d{3}[-\.\s]??\d{3}[-\.\s]??\d{4}|\(\d{3}\)\s*\d{3}[-\.\s]??\d{4}|\d{3}[-\.\s]??\d{4})'
    
    emails = re.findall(email_pattern, text)
    phones = re.findall(phone_pattern, text)
    
    return emails, phones


def extract_text_from_pdf(pdf_file):
    text = ""
    pdf_data = pdf_file.read()
    reader = fitz.Document(stream=pdf_data)
    for page in reader:
        text += page.get_text()
    return text

def extract_text_from_docx(docx_file):
    text = ""
    doc = Document(docx_file)
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    return text

def upload_cv(request):
    excel_file_url = None
    
    if request.method == 'POST':
        form = CVUploadForm(request.POST, request.FILES)
        if form.is_valid():
            cv_file = form.cleaned_data['cv_file']
            
            if cv_file.name.endswith('.pdf'):
                text = extract_text_from_pdf(cv_file)
            elif cv_file.name.endswith('.docx'):
                text = extract_text_from_docx(cv_file)
            else:
                return HttpResponse('Unsupported file format')
            
            emails, phones = extract_info(text)
            
            cv = CV(email=emails[0] if emails else None, phone=phones[0] if phones else None, text=text)
            cv.save()
            
            excel_file_url = generate_excel([cv])
            
            return render(request, 'upload_cv.html', {'form': form, 'excel_file_url': excel_file_url})
    else:
        form = CVUploadForm()
    return render(request, 'upload_cv.html', {'form': form, 'excel_file_url': excel_file_url})



#Generate excel 

def generate_excel(cvs):
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.append(['Email', 'Phone', 'Text'])
    
    for cv in cvs:
        ws.append([cv.email, cv.phone, cv.text])
    
    # Save the Excel file to the media root directory
    excel_file_path = os.path.join(settings.MEDIA_ROOT, 'cvs.xlsx')
    wb.save(excel_file_path)
    
    # Create a file storage object
    storage = FileSystemStorage()
    
    # Save the Excel file to the media storage
    with open(excel_file_path, 'rb') as excel_file:
        storage.save('cvs.xlsx', excel_file)
    
    # Construct the URL of the saved Excel file
    excel_file_url = storage.url('cvs.xlsx')
    
    return excel_file_url
